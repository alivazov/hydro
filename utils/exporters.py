"""
Модуль для экспорта данных в различные форматы
"""

import os
from datetime import datetime
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Библиотека python-docx не установлена. Экспорт в Word недоступен.")

class WordExporter:
    def __init__(self, app):
        self.app = app
        
    def export_to_word(self, file_path=None):
        """Экспорт всех данных проекта в Word"""
        if not HAS_DOCX:
            raise ImportError("Библиотека python-docx не установлена. Установите её: pip install python-docx")
            
        if not file_path:
            from tkinter import filedialog
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Документы Word", "*.docx")],
                title="Экспорт в Word"
            )
            
        if not file_path:
            return False
            
        try:
            # Создаем документ
            doc = Document()
            
            # Добавляем заголовок
            self.add_title_page(doc)
            
            # Добавляем основные расчеты
            self.add_calculations_section(doc)
            
            # Добавляем баланс
            self.add_balance_section(doc)
            
            # Добавляем таблицу площадок
            self.add_platforms_section(doc)
            
            # Добавляем этапы строительства
            self.add_construction_section(doc)
            
            # Добавляем проверку пропускной способности
            self.add_capacity_section(doc)
            
            # Сохраняем документ
            doc.save(file_path)
            return True
            
        except Exception as e:
            print(f"Ошибка экспорта в Word: {e}")
            return False
            
    def add_title_page(self, doc):
        """Добавление титульной страницы"""
        # Заголовок
        title = doc.add_heading('Гидравлический расчет канализационной сети', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о проекте
        project = self.app.project_manager.current_project
        doc.add_paragraph(f"Наименование проекта: {project.name}")
        
        if project.metadata.description:
            doc.add_paragraph(f"Описание: {project.metadata.description}")
            
        doc.add_paragraph(f"Дата создания: {project.metadata.created_date}")
        doc.add_paragraph(f"Дата расчета: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Версия: {project.metadata.version}")
        
        # Пустая строка
        doc.add_paragraph()
        
    def add_calculations_section(self, doc):
        """Добавление раздела основных расчетов"""
        doc.add_heading('1. Основные гидравлические расчеты', level=1)
        
        calculations_data = self.app.main_window.calculations_tab.get_data()
        
        if calculations_data.get('results') and calculations_data['results'] != "Результаты расчета появятся здесь...":
            doc.add_paragraph("Результаты расчетов:")
            results_paragraph = doc.add_paragraph(calculations_data['results'])
            results_paragraph.style.font.size = Pt(10)
            
        doc.add_paragraph()
        
    def add_balance_section(self, doc):
        """Добавление раздела баланса"""
        doc.add_heading('2. Баланс водопотребления и водоотведения', level=1)
        
        balance_data = self.app.main_window.balance_tab.get_data()
        if not balance_data:
            doc.add_paragraph("Данные отсутствуют")
            return
            
        # Создаем таблицу
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки таблицы
        headers = ['№', 'Обоснование', '№ пл.', '№ эт.', 'Наименование', 'Общий расход, м³/сут', '%', 'Итог %']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.runs[0].font.bold = True
            
        # Данные таблицы
        for row_data in balance_data:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                if i < len(row_cells):
                    cell = row_cells[i]
                    cell.text = str(value)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
        doc.add_paragraph()
        
    def add_platforms_section(self, doc):
        """Добавление раздела таблицы площадок"""
        doc.add_heading('3. Таблица площадок', level=1)
        
        platforms_data = self.app.main_window.platforms_tab.get_data()
        if not platforms_data:
            doc.add_paragraph("Данные отсутствуют")
            return
            
        # Создаем таблицу
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки таблицы
        headers = ['№', 'Наименование', 'Среднесуточный расход, м³/сут', 
                  'Средне-секундный расход, л/сек', 'Коэффициент неравномерности', 
                  'Секундный расчетный расход, л/сек']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.runs[0].font.bold = True
            
        # Данные таблицы
        for row_data in platforms_data:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                if i < len(row_cells):
                    cell = row_cells[i]
                    cell.text = str(value)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
        doc.add_paragraph()
        
    def add_construction_section(self, doc):
        """Добавление раздела этапов строительства"""
        doc.add_heading('4. Этапы строительства', level=1)
        
        construction_data = self.app.main_window.construction_tab.get_data()
        if not construction_data:
            doc.add_paragraph("Данные отсутствуют")
            return
            
        # Создаем таблицу
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки таблицы
        headers = ['Наименование потребителя', 'Qсут, м³/сут', 'qср, л/с']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.runs[0].font.bold = True
            
        # Данные таблицы
        for row_data in construction_data:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                if i < len(row_cells):
                    cell = row_cells[i]
                    cell.text = str(value)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
                    
        doc.add_paragraph()
        
    def add_capacity_section(self, doc):
        """Добавление раздела проверки пропускной способности"""
        doc.add_heading('5. Проверка пропускной способности', level=1)
        
        capacity_data = self.app.main_window.capacity_tab.get_data()
        if not capacity_data:
            doc.add_paragraph("Данные отсутствуют")
            return
            
        # Создаем таблицу
        table = doc.add_table(rows=1, cols=11)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки таблицы (сокращенные для удобства отображения)
        headers = ['№', 'Интервал', 'Пл./инт.', 'Qсут, м³/сут', 'Qср, л/с', 
                  'Кнеравн', 'Qрасч, л/с', 'D, мм', 'i', 'h/d', 'V, м/с']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.runs[0].font.bold = True
            
        # Данные таблицы
        for row_data in capacity_data:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                if i < len(row_cells):
                    cell = row_cells[i]
                    cell.text = str(value)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
        doc.add_paragraph()